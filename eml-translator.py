from pathlib import Path
import time
import argparse
import eml_parser
from bs4 import BeautifulSoup as bs, Stylesheet, Comment
import base64
import docx
from PyPDF2 import PdfReader
import io
import os
import magic

from openai import OpenAI
import json
import sys
from typing import Any, Dict
from urllib import request, parse, error


class LibreTranslateAPI:
    DEFAULT_URL = "https://translate.terraprint.co/"

    def __init__(self, url: str | None = None, api_key: str | None = None):
        """Create a LibreTranslate API connection.

        Args:
            url (str): The url of the LibreTranslate endpoint.
            api_key (str): The API key.
        """
        self.url = LibreTranslateAPI.DEFAULT_URL if url is None else url
        self.api_key = api_key

        # Add trailing slash
        assert len(self.url) > 0
        if self.url[-1] != "/":
            self.url += "/"

    def translate(self, q: str, source: str = "en", target: str = "es", timeout: int | None = None) -> Any:
        """Translate string

        Args:
            q (str): The text to translate
            source (str): The source language code (ISO 639)
            target (str): The target language code (ISO 639)
            timeout (int): Request timeout in seconds

        Returns:
            str: The translated text
        """
        url = self.url + "translate"
        params: Dict[str, str] = {"q": q, "source": source, "target": target}
        if self.api_key is not None:
            params["api_key"] = self.api_key
        url_params = parse.urlencode(params)
        req = request.Request(url, data=json.dumps(params).encode('utf-8'), method="POST", headers={"Content-Type": "application/json"})
        response = request.urlopen(req, timeout=timeout)
        response_str = response.read().decode()
        return json.loads(response_str)

    def detect(self, q: str, timeout: int | None = None) -> Any:
        """Detect the language of a single text.

        Args:
            q (str): Text to detect
            timeout (int): Request timeout in seconds

        Returns:
            The detected languages ex: [{"confidence": 0.6, "language": "en"}]
        """
        url = self.url + "detect"
        params: Dict[str, str] = {"q": q}
        if self.api_key is not None:
            params["api_key"] = self.api_key
        url_params = parse.urlencode(params)
        req = request.Request(url, data=url_params.encode())
        response = request.urlopen(req, timeout=timeout)
        response_str = response.read().decode()
        return json.loads(response_str)

    def languages(self, timeout: int | None = None) -> Any:
        """Retrieve list of supported languages.

        Args:
            timeout (int): Request timeout in seconds

        Returns:
            A list of available languages ex: [{"code":"en", "name":"English"}]
        """
        url = self.url + "languages"
        params: Dict[str, str] = dict()
        if self.api_key is not None:
            params["api_key"] = self.api_key
        url_params = parse.urlencode(params)
        req = request.Request(url, data=url_params.encode(), method="GET")
        response = request.urlopen(req, timeout=timeout)
        response_str = response.read().decode()
        return json.loads(response_str)


parser = argparse.ArgumentParser(
                    prog='eml-translator',
                    description='Translates EML files to english',
                    epilog='')
parser.add_argument('path')
parser.add_argument(
    '-l',
    '--language',
    help="2 letters language code.  Forces translation from this language to translate. (Defaults to auto detect)",
    required=False)
parser.add_argument(
'-s',
    '--server',
    help="https URL to the translation server",
    required=True
)
parser.add_argument(
    '-a',
    '--openaiurl',
    help="Optional.  Experimental.  When specified, this tool will mark spam emails as spam.",
    required=False
)
parser.add_argument(
    '-r',
    '--replicas',
    help="Optional.  Number of replicas of this script will run for translation.",
    required=False,
    default=1
)
parser.add_argument(
    '-i',
    '--index',
    help="Optional.  Replica index to run.  Zero based.",
    required=False,
    default=0
)
parser.add_argument(
    '-p',
    '--profile',
    help="Optional.  Used to Profile the translation performance.  It won't store results, so its easy to repeat.",
    required=False,
    default=False,
    action=argparse.BooleanOptionalAction)
args = parser.parse_args()
print("Will translate all files in " + args.path, flush=True)
translation_marker = "\n[AUTO_TRANSLATED] FROM "
profiling = args.profile

source_language = "auto"
if args.language is not None:
    source_language = args.language.lower()
target_language = "en"

lt = LibreTranslateAPI(args.server)
supported_languages = lt.languages()

ai_client = None
if args.openaiurl is not None:
    ai_client = OpenAI(base_url=args.openaiurl, api_key="lm-studio")


def numeric_hash(input):
    acc_val = 0
    for character in input:
        for byte in character.encode("utf-8"):
            acc_val += byte * 97
    return acc_val % int(args.replicas)


def replica_is_owner(input):
    hash = numeric_hash(input)
    return hash == int(args.index)


def ai_email_summarize(text):
    if ai_client is None:
        return
    completion = ai_client.chat.completions.create(
        model="QuantFactory/Meta-Llama-3-8B-Instruct-GGUF",
        messages=[
            {"role": "system", "content": "You will summarize everything I say under 256 characters."},
            {"role": "user", "content": text}
        ],
        temperature=0.7,
    )

    print(completion.choices[0].message.content, flush=True)


def save_file(file_path, data):
    if profiling:
        return

    translated_file = open(file_path, "wb")
    if isinstance(data, str):
        translated_file.write(data.encode("utf-8"))
    else:
        translated_file.write(data)
    translated_file.close()
    print(file_path + " has been saved.", flush=True)


def string_has_text(string):
    for char in string:
        if ord(char)>=65 and ord(char)<=90:
            return True
        if ord(char)>=97 and ord(char)<=122:
            return True
        if ord(char)>0xC0:
            return True


def is_noop_text(text):
    return len(text)==0 or not string_has_text(text) or is_english_charpoint(text)


def translate_text(text):
    while True:
        try:
            result = lt.translate(q=text,
                                           source=source_language, target=target_language)
            return result
        except error.HTTPError as e:
            if e.status == 500:
                return {"translatedText": text}
        except Exception as e:
            time.sleep(1)
            print("API call failed.  Retrying after 1 second.", e, flush=True)
            continue


def is_english_charpoint(string):
    for char in string:
        if not (0 < ord(char) <= 127):
            return False
    return True


def get_language_name(language_code):
    for language in supported_languages:
        if language['code'] == language_code:
            return language['name']
    return "Unknown - LibreTranslate returned no corresponding language for " + language_code


class TextBatch:
    def __init__(self):
        print("Starting translate batch.", flush=True)
        self.noop_entries = []
        self.real_entries = []
        self.batch_size = 0
        self.max_batch_size = 8192

    def add_text(self, text, context, contextParam, callback):
        if is_noop_text(text):
            self.noop_entries.append(
                dict(text=text, result=text, source_language="", context=context, contextParam=contextParam, callback=callback))
            return True
        else:
            if self.batch_size>0 and self.batch_size + len(text) >= self.max_batch_size:
                return False
            if len(text) >= self.max_batch_size:
                print("Reached max", flush=True)
            self.real_entries.append(
                dict(text=text, result=text, source_language="", context=context, contextParam=contextParam, callback=callback))
            self.batch_size += len(text)
            return True

    def finish(self):
        print("Completing translate batch of size " + str(self.batch_size) + " with " + str(len(self.noop_entries)) +
              " noop entries and " + str(len(self.real_entries)) + " real entries.", flush=True)
        if len(self.real_entries) == 0 and len(self.noop_entries) == 0:
            return
        if len(self.real_entries) > 0:
            result = translate_text([entry['text'] for entry in self.real_entries])
            translations = result["translatedText"]
            print("Batch translation completed.", flush=True)

        idx = 0
        for entry in self.real_entries:
            entry['result'] = translations[idx]
            if source_language == "auto":
                entry['source_language'] = translations[idx]["detectedLanguage"]["language"]
            else:
                entry['source_language'] = source_language
            idx = idx + 1
        for entry in self.noop_entries:
            entry['source_language'] = "en"
        for entry in self.noop_entries+self.real_entries:
            original = entry["text"]
            result = entry["result"]
            context = entry["context"]
            contextParam = entry["contextParam"]
            source_lang = entry["source_language"]
            entry["callback"](original, result, source_lang, context, contextParam)


file_count = 0


def flatten(list):
    result = ""
    for item in list:
        result += item
    return result


def docx_translated_callback(original, result, source_lang, context, contextParam):
    if source_lang != "en":
        context.text += " --- " + translation_marker + source_lang + ": " + result


def translate_docx(filename, partname, html_data):
    try:
        print("Opening docx " + filename + "-" + partname, flush=True)
        doc = docx.Document(io.BytesIO(html_data))
    except Exception:
        print("Failed to open docx " + filename + "-" + partname + " Dumping it.", flush=True)
        save_file(filename + "-" + partname, html_data)
        return
    batch = TextBatch()
    paragraph_num = len(doc.paragraphs)
    table_num = len(doc.tables)
    print("Translating " + str(paragraph_num) + ".docx paragraphs and " + str(table_num) + " tables in email " +
          filename + " attachment: " + partname, flush=True)
    currentCell = None
    visitedParagraphs = {}
    visitedCells = {}
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if id(cell) in visitedCells:
                    continue
                visitedCells[id(cell)] = True
                if cell == currentCell:
                    continue
                currentCell = cell
                for paragraph in cell.paragraphs:
                    if id(paragraph) in visitedParagraphs:
                        continue
                    visitedParagraphs[id(paragraph)] = True
                    if not batch.add_text(paragraph.text, paragraph, 0, docx_translated_callback):
                        batch.finish()
                        batch = TextBatch()
                        batch.add_text(paragraph.text, paragraph, 0, docx_translated_callback)

    paragraph_pos = 0
    for paragraph in doc.paragraphs:
        paragraph_pos += 1
        if len(paragraph.text) > 0:
            if not batch.add_text(paragraph.text, paragraph, 0, docx_translated_callback):
                batch.finish()
                batch = TextBatch()
                batch.add_text(paragraph.text, paragraph, 0, docx_translated_callback)
    batch.finish()

    if doc.inline_shapes.part is not None:
        for key in doc.inline_shapes.part.related_parts:
            related_part = doc.inline_shapes.part.related_parts[key]
            if isinstance(related_part, docx.ImagePart):
                save_file(pathStr + "-" + partname + "-" + related_part.partname.replace("/", "_"), related_part.blob)

    doc.save(pathStr + "-" + partname)


def pdf_translated_callback(original, result, source_lang, context, contextParam):
    if source_lang != "en":
        context[contextParam] = original + " --- " + translation_marker + source_lang + ": " + result;
    else:
        context[contextParam] = result;


def translate_pdf(filename, partname, pdf_data):
    batch = TextBatch()
    print("Opening PDF " + filename + "-" + partname, flush=True)
    reader = PdfReader(io.BytesIO(pdf_data))
    print("Translating " + str(len(reader.pages)) + " paragraphs in email " + filename + " attachment: " + partname, flush=True)
    output_text = ["" for _ in range(len(reader.pages))]
    index = 0
    for page in reader.pages:
        text = page.extract_text()
        if not batch.add_text(
                text,
                output_text,
                index,
                pdf_translated_callback):
            batch.finish()
            batch = TextBatch()
            batch.add_text(text,
                output_text,
                index,
                pdf_translated_callback)
        index=index+1

    batch.finish()
    return flatten(output_text)


def html_translated_callback(original, result, source_lang, context, contextParam):
    if source_lang != "en":
        context.replace_with(original + " --- " + translation_marker + source_lang + ": " + result)


def translate_html(pathStr, partName, html_data):
    batch = TextBatch()
    print("Translating HTML from email " + pathStr + " attachment: " + partName, flush=True)
    parsed_html = bs(html_data, "html.parser")
    for x in parsed_html.findAll(string=True):
        if x.string is not None and not isinstance(x.string, Comment) and not isinstance(x.string, Stylesheet):
            if not batch.add_text(
                    x.string,
                    x.string,
                    0,
                    html_translated_callback):
                batch.finish()
                batch = TextBatch()
                batch.add_text(
                    x.string,
                    x.string,
                    0,
                    html_translated_callback)

    batch.finish()
    return parsed_html.prettify(encoding='utf-8')


def translate_plain_text(pathStr, partName, data):
    print("Translating plaintext from email " + pathStr + " attachment: " + partName, flush=True)
    batch = TextBatch()
    lines = iter(data.splitlines())
    output_text = ["" for _ in range(len(data.splitlines()))]
    index = 0
    for line in lines:
        if not batch.add_text(
                line,
                output_text,
                index,
                pdf_translated_callback):
            batch.finish()
            batch = TextBatch()
            batch.add_text(
                line,
                output_text,
                index,
                pdf_translated_callback)
        index=index+1
    batch.finish()
    return flatten(output_text)


def process_email_part(contentType, pathStr, partName, data):
    if len(partName) > 100:
        toRemove = len(partName) - 100
        partName = partName[:28] + partName[28+toRemove:]

    if isinstance(data, bytes):
        if (contentType == "text/plain" or
            contentType == "application/octet-stream" or
            contentType == "text/html"):
            contentType = magic.from_buffer(data, mime=True)


    print("Processing email part " + partName + " with content-type " + contentType, flush=True)
    match contentType:
        case "text/html":
            translation = translate_html(pathStr, partName, data)
            save_file(pathStr + "-" + partName, translation)
        case "message/rfc822":
            process_eml(pathStr + "-" + partName + ".eml", data)
            save_file(pathStr + "-" + partName + ".eml", data)
        case "text/plain":
            if isinstance(data, bytes):
                save_file(pathStr + "-" + partName, data)
            else :
                translation = translate_plain_text(pathStr, partName, data)
                save_file(pathStr + "-" + partName, translation.encode("utf-8"))

        case "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            translate_docx(pathStr, partName, data)

        case "application/pdf":
            try:
                translation = translate_pdf(pathStr, partName, data)
                save_file(pathStr + "-" + partName + "-translated-content.txt", translation.encode("utf-8"))
            except Exception:
                print("Skipping translating file " + pathStr + "-" + partName, flush=True)
            save_file(pathStr + "-" + partName, data)

        case _:
            save_file(pathStr + "-" + partName, data)


def clean_eml_start(eml_bytes, max_lines):
    if max_lines == 0:
        return eml_bytes
    test = eml_bytes.split(b'\n', 1)
    if len(test[0]) > 1 and test[0][0]==ord('>'):
        return clean_eml_start(test[1], max_lines - 1)
    if b':' not in test[0]:
        return clean_eml_start(test[1], max_lines - 1)
    return eml_bytes


def process_eml(pathStr, eml_bytes):
    if len(eml_bytes) < 10:
        print("Skipping " + pathStr + " which is too short to be valid .eml file", flush=True)
        return
    eml_bytes = clean_eml_start(eml_bytes, 5)

    ep = eml_parser.EmlParser(include_attachment_data=True, include_raw_body=True)
    parsed_eml = ep.decode_email_bytes(eml_bytes)
    print("Parsed " + pathStr, flush=True)

    if "body" in parsed_eml:
        body = parsed_eml["body"]
        index = 0
        for part in body:
            index += 1
            partName = "body-" + str(index) + ".html"
            if "content_type" in part:
                content_type = part["content_type"]
                process_email_part(content_type, pathStr, partName, part["content"])

    if "attachment" in parsed_eml:
        for attachment in parsed_eml["attachment"]:
            content_hdr = attachment["content_header"]
            content_type = content_hdr["content-type"][0]
            filename = attachment["filename"]
            process_email_part(content_type, pathStr, filename, base64.b64decode(attachment["raw"]))


pathlist = Path(args.path).glob('**/*.eml')
for path in pathlist:
    file_count += 1
pathlist = Path(args.path).glob('**/*.eml')
print("Translating " + str(file_count) + " eml files", flush=True)
current_count = 0
translation_needed = False
for path in pathlist:
    print(str(current_count) + " out of " + str(file_count) + " .eml files iterated", flush=True)
    current_count += 1
    pathStr = str(path)

    if not replica_is_owner(pathStr):
        continue

    if os.path.isfile(pathStr+"-body-1.html") or os.path.isfile(pathStr+"-rtf-body.rtf"):
        print("Skipping " + pathStr + ": Already translated.", flush=True)
        # Pivot to an actual marker file to skip emails without rtf or html
        Path(pathStr+"-translated-mark.mrk").touch()
        continue

    if os.path.isfile(pathStr+"-translated-mark.mrk"):
        print("Skipping " + pathStr+": Already translated.", flush=True)
        continue

    with open(pathStr, 'rb') as fp:
        eml_bytes = raw_email = fp.read()
        print("Processing " + pathStr + " ("+str(len(eml_bytes))+" bytes)", flush=True)
        process_eml(pathStr, eml_bytes)

    if not profiling:
        Path(pathStr + "-translated-mark.mrk").touch()

print("Completed.", flush=True)